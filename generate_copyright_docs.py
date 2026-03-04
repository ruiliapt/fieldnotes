#!/usr/bin/env python3
"""
生成中国软件著作权申请材料（Word 文档）

产出：
  copyright/源代码文档.docx  — 前30页 + 后30页源代码
  copyright/软件说明书.docx  — 软件功能说明书

使用方法：
  python generate_copyright_docs.py
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── 常量 ───────────────────────────────────────────────────────────────

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "copyright")

SOFTWARE_NAME = "Fieldnotes Lite 田野笔记管理软件 V1.0"
DEVELOPER = "吕珊珊"

# 源文件拼接顺序
SOURCE_FILES = [
    "main.py",
    "database.py",
    "gui.py",
    "exporter.py",
    "theme.py",
    "ai_backend.py",
    "ai_prompts.py",
    "ai_widgets.py",
    "logger.py",
    "qt_conf_fix.py",
]

LINES_PER_PAGE = 50  # 每页代码行数
TOTAL_PAGES = 60     # 总页数（前30 + 后30）
FRONT_LINES = LINES_PER_PAGE * 30  # 前 1500 行
BACK_LINES = LINES_PER_PAGE * 30   # 后 1500 行


# ─── 工具函数 ────────────────────────────────────────────────────────────

def _set_cell_margins(section):
    """设置页边距：上下 2.54cm，左 3.17cm，右 2.54cm"""
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)


def _add_header(section, text):
    """添加页眉"""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _add_page_number_footer(section):
    """添加页脚页码（第 X 页）"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_pre = p.add_run("第 ")
    run_pre.font.size = Pt(9)
    run_pre.font.name = "宋体"
    run_pre._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # PAGE 字段
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run_page = p.add_run()
    run_page._element.append(fld_char_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run_page2 = p.add_run()
    run_page2._element.append(instr)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_page3 = p.add_run()
    run_page3._element.append(fld_char_end)

    run_suf = p.add_run(" 页")
    run_suf.font.size = Pt(9)
    run_suf.font.name = "宋体"
    run_suf._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _set_paragraph_spacing(paragraph, line_spacing_pt=None, space_after=Pt(0), space_before=Pt(0)):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before
    if line_spacing_pt:
        pf.line_spacing = line_spacing_pt


# ─── 文档一：源代码文档 ──────────────────────────────────────────────────

def _collect_source_lines():
    """读取所有源文件并拼接，每个文件开头插入分隔注释行"""
    all_lines = []
    for fname in SOURCE_FILES:
        path = os.path.join(BASE_DIR, fname)
        all_lines.append(f"# ===== 文件名: {fname} =====")
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                all_lines.append(line.rstrip("\n"))
    return all_lines


def generate_source_code_doc():
    """生成源代码文档.docx"""
    print("正在生成源代码文档...")

    all_lines = _collect_source_lines()
    total = len(all_lines)
    print(f"  源代码总行数: {total}")

    # 截取前 1500 行和后 1500 行
    front = all_lines[:FRONT_LINES]
    back = all_lines[-BACK_LINES:]

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    _set_cell_margins(section)
    _add_header(section, SOFTWARE_NAME)
    _add_page_number_footer(section)

    def _write_code_lines(lines, start_line_num=1):
        """将代码行写入文档，每 LINES_PER_PAGE 行分一页"""
        for i, line in enumerate(lines):
            line_num = start_line_num + i
            text = f"{line_num:>5}  {line}"
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, line_spacing_pt=Pt(12))
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

            # 每 LINES_PER_PAGE 行后分页（但不在最后一行后）
            if (i + 1) % LINES_PER_PAGE == 0 and (i + 1) < len(lines):
                p_last = doc.paragraphs[-1]
                run_break = p_last.add_run()
                run_break.add_break(docx.enum.text.WD_BREAK.PAGE)

    import docx.enum.text

    # 写前 30 页
    _write_code_lines(front, start_line_num=1)

    # 在前后之间插入分隔页
    p_sep = doc.add_paragraph()
    _set_paragraph_spacing(p_sep, line_spacing_pt=Pt(12))
    run_sep = p_sep.add_run("\n\n\n（此处省略中间部分源代码）\n\n\n")
    run_sep.font.size = Pt(12)
    run_sep.font.name = "宋体"
    run_sep._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run_sep.bold = True
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 分页
    run_break = p_sep.add_run()
    run_break.add_break(docx.enum.text.WD_BREAK.PAGE)

    # 写后 30 页
    back_start_num = total - BACK_LINES + 1
    _write_code_lines(back, start_line_num=back_start_num)

    out_path = os.path.join(OUTPUT_DIR, "源代码文档.docx")
    doc.save(out_path)
    print(f"  已保存: {out_path}")


# ─── 文档二：软件说明书 ──────────────────────────────────────────────────

def _add_heading(doc, text, level=1):
    """添加标题并设置中文字体"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    return h


def _add_para(doc, text, bold=False, font_size=Pt(12), font_name="宋体",
              first_line_indent=Cm(0.74), alignment=None):
    """添加正文段落"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing = Pt(22)
    if first_line_indent:
        pf.first_line_indent = first_line_indent
    if alignment:
        p.alignment = alignment

    run = p.add_run(text)
    run.font.size = font_size
    run.font.name = font_name
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    return p


def _add_bullet(doc, text, font_size=Pt(12), font_name="宋体"):
    """添加列表项"""
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(22)
    # 清除默认run，手动添加
    p.clear()
    run = p.add_run(text)
    run.font.size = font_size
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    return p


def generate_manual_doc():
    """生成软件说明书.docx"""
    print("正在生成软件说明书...")

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    _set_cell_margins(section)

    # ── 封面 ──
    for _ in range(6):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(SOFTWARE_NAME)
    run.font.size = Pt(26)
    run.font.name = "黑体"
    run.bold = True
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run("软件说明书")
    run.font.size = Pt(22)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    for _ in range(4):
        doc.add_paragraph()

    for line in [f"开发者：{DEVELOPER}", "版本：V1.0", "日期：2025年6月"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(14)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 分页
    doc.add_page_break()

    # ── 1. 引言 ──
    _add_heading(doc, "1  引言")

    _add_heading(doc, "1.1  编写目的", level=2)
    _add_para(doc,
        "本文档是 Fieldnotes Lite 田野笔记管理软件 V1.0 的软件说明书，"
        "旨在详细描述软件的功能、运行环境和操作方法，为用户提供使用指导，"
        "同时作为软件著作权登记的申请材料之一。")

    _add_heading(doc, "1.2  软件标识", level=2)
    _add_para(doc, f"软件全称：{SOFTWARE_NAME}")
    _add_para(doc, "软件简称：Fieldnotes Lite")
    _add_para(doc, "版本号：V1.0")

    _add_heading(doc, "1.3  开发者", level=2)
    _add_para(doc, f"开发者：{DEVELOPER}")

    # ── 2. 软件概述 ──
    _add_heading(doc, "2  软件概述")

    _add_heading(doc, "2.1  软件用途", level=2)
    _add_para(doc,
        "Fieldnotes Lite 是一款面向语言学田野调查研究的专业语料管理工具。"
        "软件为语言学研究人员提供了高效、直观的田野调查数据管理方案，"
        "支持语料的录入、编辑、分类、搜索、导出等核心功能。"
        "软件还集成了人工智能辅助功能，可帮助研究人员进行词汇分解和智能翻译，"
        "大幅提升田野调查的工作效率。")

    _add_heading(doc, "2.2  主要功能", level=2)
    _add_para(doc, "本软件提供以下主要功能：")
    for feat in [
        "语料数据管理：支持语料的增加、删除、修改、查询，支持分类和标签管理",
        "IPA 音标工具栏：提供国际音标（IPA）快速输入工具，涵盖元音、辅音、声调、变音符等",
        "搜索与筛选：支持全文搜索、字段搜索和高亮显示",
        "数据导出：支持导出为 Word 文档、CSV 文件和 JSON 文件",
        "AI 辅助功能：集成多种大语言模型，提供词汇分解和智能翻译功能",
        "界面主题：支持深色模式和浅色模式切换",
        "数据备份与恢复：支持数据库的自动备份和手动恢复",
        "打印功能：支持语料数据的格式化打印输出",
        "数据完整性检查：提供数据库完整性校验功能",
    ]:
        _add_bullet(doc, feat)

    _add_heading(doc, "2.3  技术架构", level=2)
    _add_para(doc,
        "本软件采用 Python 语言开发，使用 PyQt6 框架构建图形用户界面，"
        "使用 SQLite 数据库进行本地数据存储。主要技术栈如下：")
    for item in [
        "编程语言：Python 3.11+",
        "GUI 框架：PyQt6 6.x",
        "数据库：SQLite 3",
        "文档生成：python-docx",
        "数据处理：pandas",
        "AI 接口：支持 Claude API、OpenAI 兼容协议、Ollama 本地推理",
    ]:
        _add_bullet(doc, item)

    # ── 3. 运行环境 ──
    _add_heading(doc, "3  运行环境")

    _add_heading(doc, "3.1  硬件环境", level=2)
    for item in [
        "处理器：Intel / Apple Silicon 64 位处理器",
        "内存：4 GB 及以上",
        "硬盘空间：200 MB 及以上可用空间",
    ]:
        _add_bullet(doc, item)

    _add_heading(doc, "3.2  软件环境", level=2)
    for item in [
        "操作系统：macOS 10.13 及以上 / Windows 10 及以上",
        "运行时环境：Python 3.11 及以上",
        "依赖库：PyQt6、python-docx、pandas",
    ]:
        _add_bullet(doc, item)

    _add_heading(doc, "3.3  网络环境", level=2)
    _add_para(doc,
        "软件的基本功能（语料管理、搜索、导出等）无需网络连接即可使用。"
        "AI 辅助功能中的在线模式（Claude API、OpenAI 兼容 API）需要互联网连接；"
        "本地模式（Ollama）无需互联网连接。")

    # ── 4. 功能说明 ──
    _add_heading(doc, "4  功能说明")

    # 4.1
    _add_heading(doc, "4.1  语料数据管理", level=2)
    _add_para(doc,
        "语料数据管理是本软件的核心功能，提供完整的语料增删改查操作。"
        "每条语料记录包含以下字段：")
    for item in [
        "Example ID：语料编号，用于唯一标识每条语料",
        "Source Text / 原文：田野调查中记录的原始语言文本",
        "Source Text (CN) / 原文（中文）：原始语言文本的中文注释",
        "Gloss / 词注：词汇层面的逐词注释（英文）",
        "Gloss (CN) / 词注（中文）：词汇层面的逐词注释（中文）",
        "Translation / 翻译：整句翻译（英文）",
        "Translation (CN) / 翻译（中文）：整句翻译（中文）",
        "Notes / 备注：研究人员的补充说明",
        "Tags / 标签：用于分类管理的标签，多个标签用逗号分隔",
    ]:
        _add_bullet(doc, item)
    _add_para(doc,
        "用户可以通过主界面的工具栏按钮或快捷键进行语料的添加、编辑和删除操作。"
        "软件支持批量标签管理功能，可以同时为多条语料添加或移除标签。"
        "同时，软件内置语料去重功能，可检测并处理重复的语料记录。")

    # 4.2
    _add_heading(doc, "4.2  IPA 音标工具栏", level=2)
    _add_para(doc,
        "软件提供可折叠的 IPA（国际音标）符号输入工具栏，"
        "方便研究人员在录入语料时快速插入音标符号。工具栏按类别组织音标符号：")
    for item in [
        "元音：ɑ æ ɛ ə ɪ ɨ ɯ ɔ ø œ ʊ ʌ 等",
        "辅音：ŋ ɲ ɴ ʔ β ɸ θ ð ʃ ʒ ɕ ʑ ʂ ʐ 等",
        "声调：˥ ˦ ˧ ˨ ˩ 等声调标记",
        "上标：⁰ ¹ ² ³ ⁴ ⁵ ⁶ ⁷ ⁸ ⁹",
        "变音符：ʰ ʷ ʲ ʼ 等变音标记",
    ]:
        _add_bullet(doc, item)
    _add_para(doc, "用户点击任意音标符号即可将其插入到当前活动的文本输入框中。")

    # 4.3
    _add_heading(doc, "4.3  搜索与筛选", level=2)
    _add_para(doc,
        "软件提供灵活的搜索与筛选功能，帮助研究人员快速定位所需语料：")
    _add_para(doc,
        "全文搜索：在所有文本字段中搜索关键词，支持模糊匹配。"
        "搜索结果中的匹配文本会以高亮方式显示，便于用户快速识别。")
    _add_para(doc,
        "标签筛选：支持按标签筛选语料，用户可以选择一个或多个标签进行筛选，"
        "快速查看特定类别的语料。")
    _add_para(doc,
        "搜索结果支持实时更新，用户输入关键词后即时显示匹配结果。")

    # 4.4
    _add_heading(doc, "4.4  数据导出", level=2)
    _add_para(doc, "软件支持将语料数据导出为多种格式：")
    _add_para(doc,
        "Word 文档导出（.docx）：按照语言学论文的标准格式导出语料，"
        "包含编号、原文、词注、翻译等字段，支持对齐排版。"
        "导出的文档格式规范，可直接用于学术论文或研究报告。")
    _add_para(doc,
        "CSV 文件导出（.csv）：将语料数据导出为逗号分隔值文件，"
        "便于在 Excel、SPSS 等工具中进行进一步的数据分析。")
    _add_para(doc,
        "JSON 文件导出（.json）：将语料数据导出为 JSON 格式，"
        "便于与其他软件系统进行数据交换和集成。")

    # 4.5
    _add_heading(doc, "4.5  AI 辅助功能", level=2)
    _add_para(doc,
        "软件集成了人工智能辅助功能，利用大语言模型帮助研究人员进行语料分析：")
    _add_para(doc,
        "AI 词汇分解：自动对选定的语料进行词汇层面的分解分析，"
        "生成逐词注释（gloss），帮助研究人员快速完成词汇标注工作。"
        "分解结果包括词素切分、词性标注和语义解释。")
    _add_para(doc,
        "AI 智能翻译：利用大语言模型对语料进行高质量翻译，"
        "支持多种目标语言。翻译过程中考虑语言学上下文，"
        "生成更准确的学术翻译结果。")
    _add_para(doc,
        "AI 功能采用混合模式架构，支持在线模式和离线模式：")
    for item in [
        "在线模式：通过 API 调用云端大语言模型，支持 Claude、OpenAI 等",
        "离线模式：通过 Ollama 运行本地模型，无需互联网连接",
    ]:
        _add_bullet(doc, item)

    # 4.6
    _add_heading(doc, "4.6  多 LLM 支持", level=2)
    _add_para(doc,
        "软件支持多种大语言模型提供者，用户可根据需要选择合适的模型：")
    for item in [
        "Claude（Anthropic）：支持 Claude 系列模型",
        "OpenAI：支持 GPT-4o 等模型",
        "DeepSeek：支持 DeepSeek-Chat 等模型",
        "通义千问（Qwen）：支持阿里云通义千问系列模型",
        "智谱 GLM（Zhipu）：支持 GLM-4 系列模型",
        "百度文心（ERNIE）：支持文心一言系列模型",
        "Ollama（本地）：支持本地部署的开源模型",
    ]:
        _add_bullet(doc, item)
    _add_para(doc,
        "软件通过 OpenAI 兼容协议统一接入多家云端模型提供者，"
        "用户只需配置 API 密钥和选择模型即可使用。"
        "同时支持自动模式，按优先级依次尝试可用的模型提供者。")

    # 4.7
    _add_heading(doc, "4.7  界面主题", level=2)
    _add_para(doc,
        "软件支持深色模式和浅色模式两种界面主题，用户可根据个人偏好"
        "和工作环境进行切换。深色模式采用低亮度配色方案，"
        "适合在光线较暗的环境下长时间使用，有助于减轻眼部疲劳。"
        "浅色模式为默认主题，提供传统的明亮界面风格。"
        "主题切换通过菜单栏的视图菜单进行操作，切换后即时生效。")

    # 4.8
    _add_heading(doc, "4.8  数据备份与恢复", level=2)
    _add_para(doc,
        "软件提供完善的数据备份与恢复功能，确保用户数据安全：")
    _add_para(doc,
        "自动备份：软件在关键操作（如数据导入、批量修改等）前自动创建备份。"
        "备份文件存储在用户主目录的 .fieldnote/backups 目录下，"
        "文件名包含时间戳，便于识别和管理。")
    _add_para(doc,
        "手动备份：用户可通过菜单手动触发数据备份，"
        "自行选择备份文件的保存位置。")
    _add_para(doc,
        "数据恢复：当数据出现异常时，用户可以从备份文件恢复数据。"
        "恢复操作会替换当前数据库文件，恢复前软件会自动备份当前数据。")

    # 4.9
    _add_heading(doc, "4.9  打印功能", level=2)
    _add_para(doc,
        "软件支持将语料数据进行格式化打印输出。打印功能使用系统打印对话框，"
        "用户可选择打印机、设置纸张大小和打印范围。"
        "打印输出的格式与 Word 导出格式一致，"
        "包含语料编号、原文、词注、翻译等字段，排版清晰规范。")

    # 4.10
    _add_heading(doc, "4.10  数据完整性检查", level=2)
    _add_para(doc,
        "软件内置数据完整性检查功能，可对数据库进行校验，"
        "检测可能存在的数据损坏、字段缺失等问题。"
        "检查完成后向用户报告数据库状态和发现的问题，"
        "帮助用户及时发现和处理数据异常。")

    # ── 5. 操作说明 ──
    _add_heading(doc, "5  操作说明")

    # 5.1
    _add_heading(doc, "5.1  主界面布局", level=2)
    _add_para(doc, "软件主界面由以下几个部分组成：")
    for item in [
        "菜单栏：位于窗口顶部，提供文件、编辑、视图、工具、帮助等菜单",
        "工具栏：位于菜单栏下方，提供常用操作的快捷按钮（添加、编辑、删除、搜索等）",
        "IPA 工具栏：可折叠的音标符号面板，位于工具栏下方",
        "数据表格：主界面中央区域，以表格形式展示语料数据",
        "状态栏：位于窗口底部，显示当前语料数量等状态信息",
    ]:
        _add_bullet(doc, item)

    # 5.2
    LQ = "\u201c"  # "
    RQ = "\u201d"  # "
    AR = "\u2192"  # →

    _add_heading(doc, "5.2  添加语料", level=2)
    _add_para(doc, "添加新语料的操作步骤如下：")
    _add_para(doc, f"步骤一：点击工具栏中的{LQ}添加{RQ}按钮，或使用菜单{LQ}编辑{AR}添加语料{RQ}，弹出语料编辑对话框。")
    _add_para(doc, "步骤二：在对话框中填写各个字段信息，包括 Example ID、原文、词注、翻译、备注和标签等。"
              "可以使用 IPA 工具栏插入音标符号。")
    _add_para(doc, f"步骤三：填写完毕后点击{LQ}保存{RQ}按钮，语料将被添加到数据库中，并在主界面表格中显示。")

    # 5.3
    _add_heading(doc, "5.3  搜索语料", level=2)
    _add_para(doc, "搜索语料的操作步骤如下：")
    _add_para(doc, "步骤一：在工具栏的搜索框中输入关键词。")
    _add_para(doc, "步骤二：软件会实时搜索并在表格中显示匹配的语料记录，匹配的文本会以高亮方式标记。")
    _add_para(doc, f"步骤三：如需按标签筛选，可使用{LQ}标签筛选{RQ}下拉菜单选择目标标签。")
    _add_para(doc, f"步骤四：点击{LQ}清除{RQ}按钮可清除搜索条件，恢复显示所有语料。")

    # 5.4
    _add_heading(doc, "5.4  导出数据", level=2)
    _add_para(doc, "导出数据的操作步骤如下：")
    _add_para(doc, f"步骤一：选择菜单{LQ}文件{AR}导出{RQ}，在子菜单中选择导出格式（Word/CSV/JSON）。")
    _add_para(doc, "步骤二：在弹出的文件保存对话框中选择保存位置和文件名。")
    _add_para(doc, f"步骤三：点击{LQ}保存{RQ}按钮，软件将按照选定的格式导出当前显示的语料数据。")
    _add_para(doc, "步骤四：导出完成后，软件会提示导出成功信息。")

    # 5.5
    _add_heading(doc, "5.5  AI 设置与使用", level=2)
    _add_para(doc, "配置和使用 AI 辅助功能的步骤如下：")
    _add_para(doc,
        f"步骤一：选择菜单{LQ}工具{AR}AI 设置{RQ}，打开 AI 配置对话框。")
    _add_para(doc,
        "步骤二：选择 AI 提供者（Claude/OpenAI/DeepSeek/通义千问/智谱GLM/百度文心/Ollama），"
        "填写相应的 API 密钥和模型名称。如选择 Ollama，需确保本地 Ollama 服务已启动。")
    _add_para(doc,
        f"步骤三：点击{LQ}测试连接{RQ}按钮验证配置是否正确。")
    _add_para(doc,
        f"步骤四：配置完成后，在主界面选中需要分析的语料，"
        f"点击{LQ}AI 词汇分解{RQ}或{LQ}AI 翻译{RQ}按钮即可使用 AI 辅助功能。")
    _add_para(doc,
        "步骤五：AI 处理完成后，结果将自动填入对应的语料字段中。"
        "用户可以在此基础上进行手动调整和修改。")

    out_path = os.path.join(OUTPUT_DIR, "软件说明书.docx")
    doc.save(out_path)
    print(f"  已保存: {out_path}")


# ─── 主程序 ──────────────────────────────────────────────────────────────

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    generate_source_code_doc()
    generate_manual_doc()
    print("\n✅ 全部完成！文件位于 copyright/ 目录下。")


if __name__ == "__main__":
    main()
