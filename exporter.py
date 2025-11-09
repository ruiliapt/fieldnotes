"""
导出模块 - 负责将语料导出为各种格式
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Dict


class TextFormatter:
    """文本格式化类 - 生成对齐的语言学格式文本"""
    
    @staticmethod
    def calculate_display_width(text: str) -> int:
        """
        计算文本的显示宽度（考虑等宽字体）
        在等宽字体中：
        - 中文、日文等宽字符：2
        - 英文、数字、标点：1
        - 上标数字（音调）：0.5（计为1，但多个上标算一起）
        """
        import unicodedata
        
        width = 0
        i = 0
        text_len = len(text)
        
        while i < text_len:
            char = text[i]
            code = ord(char)
            
            # 中日韩统一表意文字
            if 0x4E00 <= code <= 0x9FFF:
                width += 2
            # 中日韩扩展
            elif 0x3400 <= code <= 0x4DBF:
                width += 2
            # 全角字符
            elif 0xFF00 <= code <= 0xFFEF:
                width += 2
            # 上标数字（音调标记）- 统一计为0.5宽度
            elif 0x2070 <= code <= 0x209F or 0x00B2 <= code <= 0x00B3:
                # 连续的上标数字一起计算
                superscript_count = 0
                while i < text_len and (0x2070 <= ord(text[i]) <= 0x209F or 0x00B2 <= ord(text[i]) <= 0x00B3):
                    superscript_count += 1
                    i += 1
                i -= 1  # 回退一个，因为循环会+1
                # 每2个上标字符算1个宽度
                width += (superscript_count + 1) // 2
            # 组合变音符号（紧跟在字母后面的）
            elif unicodedata.category(char) in ['Mn', 'Mc', 'Me']:
                width += 0
            # 其他字符
            else:
                width += 1
            
            i += 1
        
        return width
    
    @staticmethod
    def align_words(source_words: List[str], gloss_words: List[str]) -> tuple:
        """
        对齐原文和词汇分解（紧凑型完美左对齐）
        
        Args:
            source_words: 原文词列表
            gloss_words: 词汇分解列表
            
        Returns:
            (对齐后的原文行, 对齐后的gloss行)
        """
        # 确保两个列表长度相同
        max_len = max(len(source_words), len(gloss_words))
        source_words = source_words + [''] * (max_len - len(source_words))
        gloss_words = gloss_words + [''] * (max_len - len(gloss_words))
        
        # 计算每个位置需要的实际宽度（紧凑模式）
        column_widths = []
        for src, gls in zip(source_words, gloss_words):
            # 计算字符长度（包括上标）
            src_len = len(src)
            gls_len = len(gls)
            
            # 取较大者，并加上最少2个空格作为间距
            col_width = max(src_len, gls_len) + 2
            column_widths.append(col_width)
        
        # 对齐每一列
        aligned_source_parts = []
        aligned_gloss_parts = []
        
        for src, gls, col_width in zip(source_words, gloss_words, column_widths):
            # 左对齐，右侧填充空格到列宽
            src_padded = src.ljust(col_width)
            gls_padded = gls.ljust(col_width)
            
            aligned_source_parts.append(src_padded)
            aligned_gloss_parts.append(gls_padded)
        
        # 连接所有列
        return (''.join(aligned_source_parts), ''.join(aligned_gloss_parts))
    
    @staticmethod
    def format_entry(entry: Dict, show_numbering: bool = True, number_format: str = "()",
                    max_words_per_line: int = 10, include_chinese: bool = False) -> str:
        """
        格式化单条语料为标准语言学格式（Leipzig Glossing Rules）
        支持长句自动分行
        
        Args:
            entry: 语料记录
            show_numbering: 是否显示编号
            number_format: 编号格式 "()" 或 "."
            max_words_per_line: 每行最多显示的词数（默认10）
            include_chinese: 是否包含汉字注释字段
            
        Returns:
            格式化后的文本
        """
        lines = []
        
        # 准备编号
        numbering_text = ""
        if show_numbering:
            example_id = entry.get('example_id', '')
            if example_id:
                if number_format == "()":
                    numbering_text = f"({example_id}) "
                else:
                    numbering_text = f"{example_id}. "
            else:
                entry_id = entry.get('id', '')
                if number_format == "()":
                    numbering_text = f"({entry_id}) "
                else:
                    numbering_text = f"{entry_id}. "
        
        # 分词（按空格分割）
        source_text = entry.get('source_text', '').strip()
        gloss = entry.get('gloss', '').strip()
        translation = entry.get('translation', '').strip()
        notes = entry.get('notes', '').strip()
        
        # 汉字字段（如果包含）
        source_text_cn = entry.get('source_text_cn', '').strip() if include_chinese else ""
        gloss_cn = entry.get('gloss_cn', '').strip() if include_chinese else ""
        translation_cn = entry.get('translation_cn', '').strip() if include_chinese else ""
        
        # 计算缩进空格数（编号的长度）
        indent_spaces = ' ' * len(numbering_text)
        
        # 原文行和注释行 - 进行词对齐（支持自动分行）
        source_has_words = len(source_text.split()) > 1
        gloss_has_words = len(gloss.split()) > 1
        
        if source_has_words and gloss_has_words:
            # 多词情况：编号和原文在同一行
            source_words = source_text.split()
            gloss_words = gloss.split()
            
            # 如果原文最后一个词是单独的句号，将其合并到前一个词
            if len(source_words) > 1 and source_words[-1] == '.':
                source_words[-2] = source_words[-2] + '.'
                source_words.pop()
            
            # 确保词数相同
            max_len = max(len(source_words), len(gloss_words))
            source_words = source_words + [''] * (max_len - len(source_words))
            gloss_words = gloss_words + [''] * (max_len - len(gloss_words))
            
            # 如果词数超过每行最大值，分行显示
            if len(source_words) > max_words_per_line:
                # 计算缩进
                if numbering_text:
                    indent = ' ' * len(numbering_text)
                else:
                    indent = '    '
                
                # 分批处理
                for i in range(0, len(source_words), max_words_per_line):
                    batch_source = source_words[i:i + max_words_per_line]
                    batch_gloss = gloss_words[i:i + max_words_per_line]
                    
                    aligned_source, aligned_gloss = TextFormatter.align_words(batch_source, batch_gloss)
                    # 第一批添加编号
                    if i == 0 and numbering_text:
                        lines.append(f"{numbering_text}{aligned_source}")
                    else:
                        lines.append(f"{indent}{aligned_source}")
                    lines.append(f"{indent}{aligned_gloss}")
                    
                    # 如果还有下一批，添加空行分隔
                    if i + max_words_per_line < len(source_words):
                        lines.append("")
            else:
                # 一行能显示完
                aligned_source, aligned_gloss = TextFormatter.align_words(source_words, gloss_words)
                # 编号和原文在同一行
                if numbering_text:
                    lines.append(f"{numbering_text}{aligned_source}")
                    # 计算缩进（编号的长度）
                    indent = ' ' * len(numbering_text)
                else:
                    lines.append(f"    {aligned_source}")
                    indent = '    '
                
                # 原文(汉字) 紧跟原文
                if source_text_cn:
                    lines.append(f"{indent}【原文(汉字)】{source_text_cn}")
                
                lines.append(f"{indent}{aligned_gloss}")
                # 词汇分解(汉字) 紧跟词汇分解
                if gloss_cn:
                    lines.append(f"{indent}【词汇分解(汉字)】{gloss_cn}")
            
            # 翻译行（使用编号的缩进）
            if numbering_text:
                indent = ' ' * len(numbering_text)
            else:
                indent = '    '
            
            if translation:
                if notes:
                    lines.append(f"{indent}'{translation}' ({notes})")
                else:
                    lines.append(f"{indent}'{translation}'")
            elif notes:
                lines.append(f"{indent}({notes})")
            
            # 翻译(汉字) 紧跟翻译
            if translation_cn:
                lines.append(f"{indent}【翻译(汉字)】{translation_cn}")
        else:
            # 单词情况：编号和原文在同一行
            lines.append(f"{numbering_text}{source_text}")
            # 原文(汉字) 紧跟原文
            if source_text_cn:
                lines.append(f"{indent_spaces}【原文(汉字)】{source_text_cn}")
            
            # gloss 行（缩进对齐）
            if gloss:
                lines.append(f"{indent_spaces}{gloss}")
            # 词汇分解(汉字) 紧跟词汇分解
            if gloss_cn:
                lines.append(f"{indent_spaces}【词汇分解(汉字)】{gloss_cn}")
            
            # 翻译行（缩进对齐）
            if translation:
                if notes:
                    lines.append(f"{indent_spaces}'{translation}' ({notes})")
                else:
                    lines.append(f"{indent_spaces}'{translation}'")
            elif notes:
                lines.append(f"{indent_spaces}({notes})")
            # 翻译(汉字) 紧跟翻译
            if translation_cn:
                lines.append(f"{indent_spaces}【翻译(汉字)】{translation_cn}")
        
        return '\n'.join(lines)
    
    @staticmethod
    def format_entries(entries: List[Dict], show_numbering: bool = True, 
                      number_format: str = "()", include_chinese: bool = False) -> str:
        """
        格式化多条语料
        
        Args:
            entries: 语料记录列表
            show_numbering: 是否显示编号
            number_format: 编号格式 "()" 或 "."
            include_chinese: 是否包含汉字注释字段
            
        Returns:
            格式化后的文本
        """
        formatted = []
        for entry in entries:
            formatted.append(TextFormatter.format_entry(entry, show_numbering, number_format, 
                                                       include_chinese=include_chinese))
            formatted.append('')  # 空行分隔每条语料
        
        return '\n'.join(formatted)


class WordExporter:
    """Word文档导出类"""
    
    def __init__(self):
        """初始化导出器"""
        self.doc = Document()
    
    def _set_cell_properties(self, cell, font_size, is_content_cell=True, font_name=None, fixed_width=None, min_width=None):
        """设置单元格的XML属性和段落格式
        
        Args:
            cell: 单元格对象
            font_size: 字体大小
            is_content_cell: 是否是内容单元格（有文字的），空单元格不设置字体
            font_name: 字体名称（可选）
            fixed_width: 固定宽度（twips），如果为None则自适应
            min_width: 最小宽度（dxa），用于确保单元格足够宽
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        if fixed_width:
            tcW.set(qn('w:w'), str(fixed_width))
            tcW.set(qn('w:type'), 'dxa')
        elif min_width:
            tcW.set(qn('w:w'), str(min_width))
            tcW.set(qn('w:type'), 'dxa')
        else:
            tcW.set(qn('w:w'), '0')
            tcW.set(qn('w:type'), 'auto')
        tcPr.append(tcW)
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'top')
        tcPr.append(vAlign)
        
        # 设置单元格不换行，防止单词内部被拆分
        noWrap = OxmlElement('w:noWrap')
        tcPr.append(noWrap)
        
        tcMar = OxmlElement('w:tcMar')
        # 左右边距设置为0.19cm (108 dxa)
        for side in ['left', 'right']:
            margin = OxmlElement(f'w:{side}')
            margin.set(qn('w:w'), '108')  # 0.19cm = 108 dxa
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        # 上下边距保持为0
        for side in ['top', 'bottom']:
            margin = OxmlElement(f'w:{side}')
            margin.set(qn('w:w'), '0')
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        tcPr.append(tcMar)
        
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            
            # 段落级别禁止换行
            pPr = paragraph._element.get_or_add_pPr()
            keepLines = OxmlElement('w:keepLines')
            pPr.append(keepLines)
            
            if is_content_cell:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    if font_name and font_name != "系统默认":
                        run.font.name = font_name
    
    def _estimate_word_height(self, word: str, font_size: int = 10) -> int:
        """估算词在单元格中的显示高度（twips）
        
        考虑因素：
        - 基础字体大小
        - 上标/下标字符（增加高度）
        - 组合符号
        
        Args:
            word: 单词字符串
            font_size: 字体大小（点）
            
        Returns:
            估算的高度（twips），1点 = 20 twips
        """
        if not word:
            return font_size * 20
        
        # 基础高度（字体大小转twips）
        base_height = font_size * 20
        
        # 检查是否有上标或下标
        has_superscript = any(char in '⁰¹²³⁴⁵⁶⁷⁸⁹' for char in word)
        has_subscript = any(char in '₀₁₂₃₄₅₆₇₈₉' for char in word)
        
        # 上标或下标会增加高度需求
        if has_superscript or has_subscript:
            base_height = int(base_height * 1.4)  # 增加40%
        
        # 添加上下边距空间（基于字体大小动态计算）
        padding = int(font_size * 4)  # 上下各留字体大小的约20%空间
        
        return base_height + padding
    
    def _estimate_word_width(self, word: str) -> float:
        """估算词在表格单元格中的显示宽度（包含边距和渲染空间）
        
        Args:
            word: 单词字符串
            
        Returns:
            估算的显示宽度（相对单位）
        """
        if not word:
            return 0
        
        # 基础字符宽度
        char_width = len(word) * 1.2
        
        # 考虑单元格左右边距（0.19cm * 2）和一些渲染缓冲
        margin_buffer = 4.0
        
        return char_width + margin_buffer
    
    def _will_cause_overflow(self, current_line_words: List[str], new_word: str, font_size: int = 10) -> bool:
        """检测添加新词后是否会导致单元格压缩或超出页宽
        
        策略：
        - 估算当前行所有词加上新词的总宽度
        - 估算每个单元格需要的最小宽度（词宽 + 边距）
        - 计算表格需要的总宽度（包括编号列）
        - 与页面可用宽度比较
        
        Args:
            current_line_words: 当前行已有的词列表
            new_word: 要添加的新词
            
        Returns:
            True: 会导致压缩或溢出，应该换行
            False: 可以安全添加
        """
        # 计算添加新词后这一行所有词的宽度
        all_words = current_line_words + [new_word]
        
        # 每个词的单元格需要的宽度（词宽 + 左右边距0.19cm*2 ≈ 216 dxa）
        # 转换为相对单位：0.19cm ≈ 108 dxa，两边 = 216 dxa ≈ 相对单位3.8
        cell_overhead = 3.8
        
        total_width = sum(self._estimate_word_width(w) + cell_overhead for w in all_words)
        
        # 加上编号列的宽度（估算为一个词的平均宽度）
        numbering_col_width = 10.0 if all_words else 0
        total_width += numbering_col_width
        
        # 动态计算Word A4页面可用宽度（基于字体大小）
        # Word A4可用宽度：15.9cm（21cm - 2.54cm左边距 - 2.54cm右边距）
        # 根据字体大小计算每cm可以容纳的字符数
        # 经验公式：10pt字体约4个字符/cm
        chars_per_cm = 40.0 / font_size  # 字体越大，每cm字符数越少
        total_chars = 15.9 * chars_per_cm  # 15.9cm可容纳的字符数
        
        # 转换为我们的相对单位（1字符 ≈ 1.2相对单位）
        base_conversion = 1.2
        
        # 动态校准系数：根据字体大小调整
        # 基准：10pt字体使用1.2的校准系数
        # 字体越大，需要的校准系数越大（因为字符间距增加）
        # 字体越小，校准系数越小
        base_font_size = 10.0
        base_calibration = 1.2
        calibration_factor = base_calibration * (font_size / base_font_size)
        
        page_available_width = total_chars * base_conversion * calibration_factor
        
        return total_width > page_available_width
    
    def _split_words_by_cumulative_width(self, words: List[str], font_size: int = 10) -> List[List[str]]:
        """完全动态分行：逐词检测是否会导致压缩或溢出
        
        策略：
        - 逐个词尝试添加到当前行
        - 检测添加后是否会导致单元格被压缩或超出页宽
        - 如果会，则换行
        - 完全基于实际检测，无预设值
        - 根据字体大小动态计算页面可用宽度
        
        Args:
            words: 词列表
            font_size: 字体大小（点），用于动态计算页面容量
            
        Returns:
            分行后的词列表，每个子列表代表一行
        """
        if not words:
            return [[]]
        
        lines = []
        current_line = []
        
        for word in words:
            # 如果当前行为空，直接添加第一个词
            if not current_line:
                current_line.append(word)
            # 检测添加这个词是否会导致压缩或溢出
            elif self._will_cause_overflow(current_line, word, font_size):
                # 会导致问题，换行
                lines.append(current_line)
                current_line = [word]
            else:
                # 不会导致问题，继续添加
                current_line.append(word)
        
        # 添加最后一行
        if current_line:
            lines.append(current_line)
        
        return lines if lines else [[]]
    
    def export(self, entries: List[Dict], output_path: str, 
               table_width: float = 5.0,
               font_size: int = 10,
               line_spacing: float = 1.15,
               show_numbering: bool = True,
               entries_per_page: int = 10,
               include_chinese: bool = False,
               max_words_per_line: int = 10,
               font_config: Dict = None) -> bool:
        """
        导出语料到Word文档（词对词对齐，透明表格）
        
        Args:
            entries: 语料记录列表
            output_path: 输出文件路径
            table_width: 表格宽度（英寸）
            font_size: 字体大小（磅）
            line_spacing: 行距
            show_numbering: 是否显示编号
            entries_per_page: 每页语料数（暂未实现分页）
            include_chinese: 是否包含汉字注释字段
            font_config: 字体配置字典，包含各字段的字体名称和大小
            
        Returns:
            是否导出成功
        """
        try:
            # 初始化字体配置
            if font_config is None:
                font_config = {}
            
            # 设置默认字体配置
            source_font = font_config.get("source_text", "Doulos SIL Compact")
            source_size = font_config.get("source_text_size", 12)
            gloss_font = font_config.get("gloss", "Charis SIL Compact")
            gloss_size = font_config.get("gloss_size", 11)
            translation_font = font_config.get("translation", None)
            translation_size = font_config.get("translation_size", 11)
            chinese_font = font_config.get("chinese", None)
            chinese_size = font_config.get("chinese_size", 10)
            
            self.doc = Document()
            
            for idx, entry in enumerate(entries, 1):
                # 准备编号（如果需要）
                numbering_text = ""
                if show_numbering:
                    example_id = entry.get('example_id', '')
                    if example_id:
                        numbering_text = f"({example_id}) "
                    else:
                        numbering_text = f"({idx}) "
                
                # 分词
                source_text = entry.get('source_text', '').strip()
                gloss = entry.get('gloss', '').strip()
                translation = entry.get('translation', '').strip()
                notes = entry.get('notes', '').strip()
                
                # 汉字字段（如果包含）
                source_text_cn = entry.get('source_text_cn', '').strip() if include_chinese else ""
                gloss_cn = entry.get('gloss_cn', '').strip() if include_chinese else ""
                translation_cn = entry.get('translation_cn', '').strip() if include_chinese else ""
                
                # 如果有多个词（分词后长度>1），进行词对齐
                source_has_words = ' ' in source_text and len(source_text.split()) > 1
                gloss_has_words = ' ' in gloss and len(gloss.split()) > 1
                if source_has_words and gloss_has_words:
                    source_words = source_text.split()
                    gloss_words = gloss.split()
                    
                    # 如果原文最后一个词是单独的句号，将其合并到前一个词
                    if len(source_words) > 1 and source_words[-1] == '.':
                        source_words[-2] = source_words[-2] + '.'
                        source_words.pop()
                    
                    # 分词处理汉字字段
                    source_words_cn = source_text_cn.split() if (include_chinese and source_text_cn) else []
                    gloss_words_cn = gloss_cn.split() if (include_chinese and gloss_cn) else []
                    
                    # 确保原文和词汇分解词数相同
                    max_len = max(len(source_words), len(gloss_words))
                    source_words = source_words + [''] * (max_len - len(source_words))
                    gloss_words = gloss_words + [''] * (max_len - len(gloss_words))
                    
                    # 完全动态分行：逐词累计宽度，超过阈值时自动换行
                    # 不使用硬编码词数，完全基于实际宽度判断
                    # 传入字体大小，动态计算页面可用宽度
                    source_lines_list = self._split_words_by_cumulative_width(source_words, source_size)
                    
                    # 根据原文的分行方式，同步分gloss和汉字字段
                    gloss_lines_list = []
                    source_cn_lines_list = []
                    gloss_cn_lines_list = []
                    
                    word_idx = 0
                    for source_line in source_lines_list:
                        line_len = len(source_line)
                        
                        # gloss按相同索引分行
                        gloss_lines_list.append(gloss_words[word_idx:word_idx + line_len])
                        
                        # 汉字字段也按相同方式分行
                        if source_words_cn and word_idx + line_len <= len(source_words_cn):
                            source_cn_lines_list.append(source_words_cn[word_idx:word_idx + line_len])
                        if gloss_words_cn and word_idx + line_len <= len(gloss_words_cn):
                            gloss_cn_lines_list.append(gloss_words_cn[word_idx:word_idx + line_len])
                        
                        word_idx += line_len
                    
                    # 计算表格尺寸
                    source_line_count = len(source_lines_list)
                    gloss_line_count = len(gloss_lines_list)
                    
                    # 计算总行数：原文 + 原文(汉字) + gloss + gloss(汉字) + 翻译 + 翻译(汉字)
                    total_rows = source_line_count + gloss_line_count + 1  # 原文 + gloss + 翻译
                    if include_chinese:
                        if source_words_cn:
                            total_rows += len(source_cn_lines_list)  # 原文(汉字)行数与原文相同
                        if gloss_words_cn:
                            total_rows += len(gloss_cn_lines_list)  # 词汇分解(汉字)行数与gloss相同
                        if translation_cn:
                            total_rows += 1  # 翻译(汉字)单行合并
                    
                    # 动态计算所有词的最大高度
                    all_words = []
                    all_words.extend(source_words)
                    all_words.extend(gloss_words)
                    if source_words_cn:
                        all_words.extend(source_words_cn)
                    if gloss_words_cn:
                        all_words.extend(gloss_words_cn)
                    
                    # 计算每个词的高度，找到最大值
                    max_height = 0
                    for word in all_words:
                        if word:  # 跳过空字符串
                            word_height = self._estimate_word_height(word, source_size)
                            max_height = max(max_height, word_height)
                    
                    # 如果没有找到有效高度，使用默认值
                    if max_height == 0:
                        max_height = 300  # 默认最小高度
                    
                    # 找出所有行中最长的一行（词数最多）
                    all_line_lists = [source_lines_list, gloss_lines_list]
                    if source_cn_lines_list:
                        all_line_lists.append(source_cn_lines_list)
                    if gloss_cn_lines_list:
                        all_line_lists.append(gloss_cn_lines_list)
                    
                    max_cols_in_line = max(
                        max(len(line) for line in line_list)
                        for line_list in all_line_lists
                    )
                    table_cols = max_cols_in_line + 1  # +1 for numbering column
                    
                    table = self.doc.add_table(rows=total_rows, cols=table_cols)
                    
                    # 设置表格属性
                    tbl = table._tbl
                    tblPr = tbl.tblPr
                    if tblPr is None:
                        tblPr = OxmlElement('w:tblPr')
                        tbl.insert(0, tblPr)
                    
                    # 设置无边框
                    tblBorders = OxmlElement('w:tblBorders')
                    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                        border = OxmlElement(f'w:{border_name}')
                        border.set(qn('w:val'), 'none')
                        border.set(qn('w:sz'), '0')
                        border.set(qn('w:space'), '0')
                        border.set(qn('w:color'), 'auto')
                        tblBorders.append(border)
                    tblPr.append(tblBorders)
                    
                    # 表格布局：自适应，根据内容和页面宽度自动调整
                    tblLayout = OxmlElement('w:tblLayout')
                    tblLayout.set(qn('w:type'), 'auto')
                    tblPr.append(tblLayout)
                    
                    # 表格宽度：占满页面宽度（100%）
                    tblW = OxmlElement('w:tblW')
                    tblW.set(qn('w:w'), '5000')  # 5000 = 100%
                    tblW.set(qn('w:type'), 'pct')  # 百分比类型
                    tblPr.append(tblW)
                    
                    # 单元格间距：0
                    tblCellSpacing = OxmlElement('w:tblCellSpacing')
                    tblCellSpacing.set(qn('w:w'), '0')
                    tblCellSpacing.set(qn('w:type'), 'dxa')
                    tblPr.append(tblCellSpacing)
                    
                    # 单元格边距
                    tblCellMar = OxmlElement('w:tblCellMar')
                    for side in ['left', 'right']:
                        margin = OxmlElement(f'w:{side}')
                        margin.set(qn('w:w'), '108')  # 0.19cm = 108 dxa
                        margin.set(qn('w:type'), 'dxa')
                        tblCellMar.append(margin)
                    for side in ['top', 'bottom']:
                        margin = OxmlElement(f'w:{side}')
                        margin.set(qn('w:w'), '0')
                        margin.set(qn('w:type'), 'dxa')
                        tblCellMar.append(margin)
                    tblPr.append(tblCellMar)
                    
                    # 设置所有行统一的精确行高，使用动态计算的最大高度
                    for row in table.rows:
                        tr = row._tr
                        trPr = tr.get_or_add_trPr()
                        trHeight = trPr.find(qn('w:trHeight'))
                        if trHeight is not None:
                            trPr.remove(trHeight)
                        trHeight = OxmlElement('w:trHeight')
                        # 使用动态计算的最大词高度，强制所有行高度完全一致
                        trHeight.set(qn('w:val'), str(max_height))
                        trHeight.set(qn('w:hRule'), 'exact')
                        trPr.append(trHeight)
                    
                    # 关键：清空所有行的所有单元格
                    for row_idx in range(total_rows):
                        for cell in table.rows[row_idx].cells:
                            cell._element.clear_content()
                            cell.text = ''  # 添加空段落
                    
                    # 追踪当前填充到哪一行
                    current_row = 0
                    
                    # 交错填充：原文行和词汇分解行紧挨着
                    # 这样可以保持每个词的上下对齐关系更直观
                    for line_idx in range(len(source_lines_list)):
                        # 1. 填充原文行
                        source_line_words = source_lines_list[line_idx]
                        
                        # 第0列：第一行放编号，其他行留空
                        cell = table.rows[current_row].cells[0]
                        if line_idx == 0 and numbering_text:
                            cell.text = numbering_text
                            self._set_cell_properties(cell, source_size, is_content_cell=True, font_name=source_font)
                        else:
                            cell.text = ''
                            self._set_cell_properties(cell, source_size, is_content_cell=False)
                        
                        # 第1到N列：填充这一行的词
                        for col_idx, word in enumerate(source_line_words):
                            cell = table.rows[current_row].cells[col_idx + 1]
                            cell.text = word
                            self._set_cell_properties(cell, source_size, is_content_cell=True, font_name=source_font)
                        
                        current_row += 1
                        
                        # 2. 填充原文(汉字)行（如果有）
                        if line_idx < len(source_cn_lines_list):
                            source_cn_line_words = source_cn_lines_list[line_idx]
                            
                            # 第0列：留空
                            cell = table.rows[current_row].cells[0]
                            cell.text = ''
                            self._set_cell_properties(cell, chinese_size, is_content_cell=False)
                            
                            # 第1到N列：填充这一行的汉字词
                            for col_idx, word in enumerate(source_cn_line_words):
                                cell = table.rows[current_row].cells[col_idx + 1]
                                cell.text = word
                                self._set_cell_properties(cell, chinese_size, is_content_cell=True, font_name=chinese_font)
                            
                            current_row += 1
                        
                        # 3. 填充gloss行
                        if line_idx < len(gloss_lines_list):
                            gloss_line_words = gloss_lines_list[line_idx]
                            
                            # 第0列：留空
                            cell = table.rows[current_row].cells[0]
                            cell.text = ''
                            self._set_cell_properties(cell, gloss_size, is_content_cell=False)
                            
                            # 第1到N列：填充这一行的gloss词
                            for col_idx, word in enumerate(gloss_line_words):
                                cell = table.rows[current_row].cells[col_idx + 1]
                                cell.text = word
                                self._set_cell_properties(cell, gloss_size, is_content_cell=True, font_name=gloss_font)
                            
                            current_row += 1
                        
                        # 4. 填充词汇分解(汉字)行（如果有）
                        if line_idx < len(gloss_cn_lines_list):
                            gloss_cn_line_words = gloss_cn_lines_list[line_idx]
                            
                            # 第0列：留空
                            cell = table.rows[current_row].cells[0]
                            cell.text = ''
                            self._set_cell_properties(cell, chinese_size, is_content_cell=False)
                            
                            # 第1到N列：填充这一行的汉字词
                            for col_idx, word in enumerate(gloss_cn_line_words):
                                cell = table.rows[current_row].cells[col_idx + 1]
                                cell.text = word
                                self._set_cell_properties(cell, chinese_size, is_content_cell=True, font_name=chinese_font)
                            
                            current_row += 1
                    
                    # 填充翻译行（合并所有列）
                    # 第0列：留空
                    cell = table.rows[current_row].cells[0]
                    cell.text = ''
                    self._set_cell_properties(cell, translation_size, is_content_cell=False)
                    
                    if translation:
                        # 合并第1列到最后一列
                        merged_cell = table.rows[current_row].cells[1]
                        for col_idx in range(2, table_cols):
                            merged_cell.merge(table.rows[current_row].cells[col_idx])
                        merged_cell.text = translation
                        self._set_cell_properties(merged_cell, translation_size, is_content_cell=True, font_name=translation_font)
                    
                    current_row += 1
                    
                    # 填充翻译(汉字)行（如果有）
                    if include_chinese and translation_cn:
                        # 第0列：留空
                        cell = table.rows[current_row].cells[0]
                        cell.text = ''
                        self._set_cell_properties(cell, chinese_size, is_content_cell=False)
                        
                        # 合并第1列到最后一列
                        merged_cell = table.rows[current_row].cells[1]
                        for col_idx in range(2, table_cols):
                            merged_cell.merge(table.rows[current_row].cells[col_idx])
                        merged_cell.text = translation_cn
                        self._set_cell_properties(merged_cell, chinese_size, is_content_cell=True, font_name=chinese_font)
                    
                    # 汉字字段已经集成到表格中
                    
                else:
                    # 如果没有分词，使用段落格式（不用表格）
                    from docx.shared import Cm
                    
                    # 计算缩进量（编号的宽度）
                    # 假设每个字符约 0.15cm，根据编号长度计算
                    indent_cm = len(numbering_text) * 0.15 if numbering_text else 0
                    
                    # 第一行：编号 + 原文
                    p1 = self.doc.add_paragraph()
                    if numbering_text:
                        p1.add_run(numbering_text).font.size = Pt(font_size)
                    p1.add_run(source_text).font.size = Pt(font_size)
                    p1.paragraph_format.space_before = Pt(0)
                    p1.paragraph_format.space_after = Pt(0)
                    p1.paragraph_format.line_spacing = 1.0
                    
                    # 原文(汉字) 紧跟原文（只在include_chinese时显示）
                    if include_chinese and source_text_cn:
                        p_cn = self.doc.add_paragraph(source_text_cn)
                        p_cn.runs[0].font.size = Pt(font_size - 1)
                        p_cn.paragraph_format.left_indent = Cm(indent_cm)
                        p_cn.paragraph_format.space_before = Pt(0)
                        p_cn.paragraph_format.space_after = Pt(0)
                        p_cn.paragraph_format.line_spacing = 1.0
                    
                    # 第二行：gloss（左缩进与原文对齐）
                    if gloss:
                        p2 = self.doc.add_paragraph(gloss)
                        p2.runs[0].font.size = Pt(font_size)
                        p2.paragraph_format.left_indent = Cm(indent_cm)
                        p2.paragraph_format.space_before = Pt(0)
                        p2.paragraph_format.space_after = Pt(0)
                        p2.paragraph_format.line_spacing = 1.0
                    
                    # 词汇分解(汉字) 紧跟词汇分解（只在include_chinese时显示）
                    if include_chinese and gloss_cn:
                        p_cn = self.doc.add_paragraph(gloss_cn)
                        p_cn.runs[0].font.size = Pt(font_size - 1)
                        p_cn.paragraph_format.left_indent = Cm(indent_cm)
                        p_cn.paragraph_format.space_before = Pt(0)
                        p_cn.paragraph_format.space_after = Pt(0)
                        p_cn.paragraph_format.line_spacing = 1.0
                    
                    # 第三行：翻译（左缩进与原文对齐）
                    if translation:
                        p3 = self.doc.add_paragraph(f"'{translation}'")
                        p3.runs[0].font.size = Pt(font_size)
                        p3.paragraph_format.left_indent = Cm(indent_cm)
                        p3.paragraph_format.space_before = Pt(0)
                        p3.paragraph_format.space_after = Pt(0)
                        p3.paragraph_format.line_spacing = 1.0
                    
                    # 翻译(汉字) 紧跟翻译（只在include_chinese时显示）
                    if include_chinese and translation_cn:
                        p_cn = self.doc.add_paragraph(f"'{translation_cn}'")
                        p_cn.runs[0].font.size = Pt(font_size - 1)
                        p_cn.paragraph_format.left_indent = Cm(indent_cm)
                        p_cn.paragraph_format.space_before = Pt(0)
                        p_cn.paragraph_format.space_after = Pt(0)
                        p_cn.paragraph_format.line_spacing = 1.0
                
                # 添加备注（在表格外，与原文第一个词对齐）
                if notes:
                    note_p = self.doc.add_paragraph(f"    ({notes})")
                    note_p.runs[0].font.size = Pt(font_size - 1)
                    note_p.runs[0].italic = True
                    note_p.paragraph_format.space_before = Pt(0)
                    note_p.paragraph_format.space_after = Pt(0)
                    note_p.paragraph_format.line_spacing = 1.0
                
                # 添加段落间距
                self.doc.add_paragraph()
            
            # 保存文档
            self.doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"导出失败: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def export_custom_format(self, entries: List[Dict], output_path: str,
                            format_template: str = None) -> bool:
        """
        使用自定义格式导出
        
        Args:
            entries: 语料记录列表
            output_path: 输出文件路径
            format_template: 自定义格式模板（未来功能）
            
        Returns:
            是否导出成功
        """
        # 未来可扩展的自定义格式导出
        return self.export(entries, output_path)

